#!/usr/bin/env python3
"""Generate a formatted .docx newsletter from the markdown output."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
import os
import re

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
md_path = sys.argv[1] if len(sys.argv) > 1 else "output/2026-03-22.md"
output_path = sys.argv[2] if len(sys.argv) > 2 else os.path.expanduser("~/Desktop/Catchup with Claude/catchup-with-claude.docx")
LOGO_PATH = os.path.join(SCRIPT_DIR, "assets", "claude-logo.jpg")

# Read the markdown
with open(md_path) as f:
    raw = f.read()

# Strip markdown artifacts
raw = raw.replace("Here's the newsletter:", "").strip()
# Remove trailing meta-comments from Claude
raw = re.sub(r'\n(The file write got blocked|Want me to try|Is copying from above).*$', '', raw, flags=re.DOTALL)
# Strip markdown bold markers
raw = re.sub(r'\*\*([^*]+)\*\*', r'\1', raw)
# Strip markdown italic
raw = re.sub(r'\*([^*]+)\*', r'\1', raw)
# Strip leading ---
raw = raw.strip().strip('-').strip()

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

def add_hyperlink(paragraph, text, url, font_size=Pt(10.5), color=RGBColor(0x1a, 0x6b, 0xb5)):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), str(color).replace('#', ''))
    rPr.append(c)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_text_with_links(paragraph, text, font_size=Pt(10.5), font_color=RGBColor(0x33, 0x33, 0x33)):
    """Parse text for markdown links [text](url) and render as mixed runs + hyperlinks."""
    link_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    last_end = 0
    for match in link_pattern.finditer(text):
        # Add plain text before the link
        before = text[last_end:match.start()]
        if before:
            run = paragraph.add_run(before)
            run.font.size = font_size
            run.font.color.rgb = font_color
        # Add the hyperlink
        add_hyperlink(paragraph, match.group(1), match.group(2), font_size=font_size)
        last_end = match.end()
    # Add remaining plain text
    remaining = text[last_end:]
    if remaining:
        run = paragraph.add_run(remaining)
        run.font.size = font_size
        run.font.color.rgb = font_color

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '6',
        qn('w:space'): '1', qn('w:color'): 'CCCCCC'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

# Parse sections from the markdown
lines = raw.split('\n')
i = 0

# Logo
if os.path.exists(LOGO_PATH):
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_after = Pt(6)
    run = logo_p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(1.5))

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(2)
run = title_p.add_run("Catchup with Claude")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

# Find date line
date_line = ""
tagline_line = ""
for line in lines[:10]:
    line_clean = line.strip()
    if re.match(r'(Anthropic Weekly|Catchup with Claude).*March|March.*\d{4}', line_clean, re.IGNORECASE):
        # Extract just the date part
        date_match = re.search(r'(March\s+\d+.*?\d{4})', line_clean)
        if date_match:
            date_line = date_match.group(1)
    if 'weekly briefing' in line_clean.lower():
        tagline_line = line_clean

if date_line:
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    run = sub.add_run(date_line)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

if tagline_line:
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.paragraph_format.space_after = Pt(4)
    run = tag.add_run(tagline_line)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_hr(doc)

# Section headers to detect
section_headers = [
    "UPDATES TO CLAUDE",
    "ANTHROPIC AND AI NEWS",
    "COMMUNITY AND INNOVATION SPOTLIGHT",
    "QUICK LINKS",
    "SUGGESTED NEW FOLLOWS",
    "Curated weekly",
]

def is_section_header(line):
    clean = line.strip().strip('#').strip()
    for h in section_headers:
        if clean.upper() == h.upper() or clean.upper().startswith(h.upper()):
            return clean
    return None

# Process content
current_section = None
paragraph_buffer = []

def add_item_title(doc, text):
    """Add a bold item title."""
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

def flush_paragraph(doc, text, is_quick_link=False):
    if not text.strip():
        return
    text = text.strip()

    if is_quick_link:
        # Remove leading dashes
        text = re.sub(r'^--\s*', '', text).strip()
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_text_with_links(p, text, font_size=Pt(10), font_color=RGBColor(0x44, 0x44, 0x44))
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        add_text_with_links(p, text, font_size=Pt(10.5), font_color=RGBColor(0x33, 0x33, 0x33))

# Skip header lines we already processed
skip_until_content = True

for line in lines:
    stripped = line.strip()

    # Skip the title/tagline area
    if skip_until_content:
        if is_section_header(stripped):
            skip_until_content = False
        else:
            continue

    # Check for section header
    header = is_section_header(stripped)
    if header:
        # Flush any buffered paragraph
        if paragraph_buffer:
            is_ql = current_section and "QUICK" in current_section.upper()
            flush_paragraph(doc, ' '.join(paragraph_buffer), is_ql)
            paragraph_buffer = []

        if "curated weekly" in header.lower():
            add_hr(doc)
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.paragraph_format.space_before = Pt(8)
            run = footer.add_run(header)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            current_section = "footer"
            continue

        if current_section:
            add_hr(doc)

        current_section = header
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        continue

    if current_section == "footer":
        continue

    # Handle --- as separator
    if stripped == '---':
        if paragraph_buffer:
            is_ql = current_section and "QUICK" in current_section.upper()
            flush_paragraph(doc, ' '.join(paragraph_buffer), is_ql)
            paragraph_buffer = []
        continue

    # Empty line = paragraph break
    if not stripped:
        if paragraph_buffer:
            is_ql = current_section and "QUICK" in current_section.upper()
            flush_paragraph(doc, ' '.join(paragraph_buffer), is_ql)
            paragraph_buffer = []
        continue

    # Detect item titles: short lines (under 80 chars) that stand alone
    # between blank lines, in a content section (not quick links or footer)
    is_content_section = current_section and "QUICK" not in current_section.upper() and "FOLLOW" not in current_section.upper() and current_section != "footer"
    if is_content_section and not paragraph_buffer and len(stripped) < 80 and not stripped.startswith('--'):
        # Peek ahead: if next non-empty line is longer, this is a title
        remaining = lines[lines.index(line)+1:] if line in lines else []
        next_content = ""
        for future_line in remaining:
            if future_line.strip():
                next_content = future_line.strip()
                break
        if next_content and len(next_content) > 80:
            add_item_title(doc, stripped)
            continue

    # Accumulate paragraph
    paragraph_buffer.append(stripped)

# Flush remaining
if paragraph_buffer:
    is_ql = current_section and "QUICK" in current_section.upper()
    flush_paragraph(doc, ' '.join(paragraph_buffer), is_ql)

doc.save(output_path)
print(f"Saved to: {output_path}")
